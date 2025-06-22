import os
from PyPDF2 import PdfReader
from docx import Document

def load_text_from_file(filepath):
    if filepath.endswith(".pdf"):
        reader = PdfReader(filepath)
        return "\n".join(page.extract_text() for page in reader.pages)
    elif filepath.endswith(".docx"):
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)
    elif filepath.endswith(".txt"):
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    return ""

def get_all_sop_chunks(folder, chunk_size=500):
    
    from langchain.docstore.document import Document

    docs = []
    for file in os.listdir(folder):
        path = os.path.join(folder, file)
        content = load_text_from_file(path)
        chunks = [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
        docs.extend([Document(page_content=chunk) for chunk in chunks])
    return docs
